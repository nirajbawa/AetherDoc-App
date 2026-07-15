import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from core.config import settings

# --- Color Constants (Navy/Blue Theme) ---
PRIMARY_COLOR = RGBColor(30, 58, 138)     # Navy (#1E3A8A)
SECONDARY_COLOR = RGBColor(59, 130, 246)  # Medium Blue (#3B82F6)
TEXT_COLOR = RGBColor(31, 41, 55)         # Dark Grey (#1F2937)
BORDER_COLOR_HEX = "1E3A8A"
CALLOUT_BG_HEX = "EFF6FF"
SHADING_HEX_ALT = "F9FAFB"               # Soft grey for alternating table rows
SHADING_HEX_HEADER = "1E3A8A"            # Navy for table header

def set_cell_background(cell, fill_hex):
    """Sets the background color (shading) of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_left_border(cell, color_hex="1E3A8A", size_pt=3):
    """Sets a thick left border for a cell and removes top/bottom/right borders (for Callouts)."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:left w:val="single" w:sz="{int(size_pt * 8)}" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="none"/>
            <w:right w:val="none"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def add_page_number_to_run(run):
    """Appends dynamic page number XML field to a run."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def style_heading_paragraph(p, text, size_pt, color, bold=True, space_before=12, space_after=6):
    """Applies specific heading styling to a paragraph."""
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.color.rgb = color
    return p

def compile_document_to_docx(session_id: str, title: str, sections: list) -> str:
    """
    Compiles a structured list of sections into a styled .docx document.
    Each section is a dictionary: {"heading": "Section Heading", "content": "Text content"}
    Returns the absolute path to the generated document.
    """
    doc = Document()
    
    # Configure document margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Configure Header
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header_p.add_run(f"Document Generation Agent  |  {title}")
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(8.5)
        header_run.font.color.rgb = RGBColor(156, 163, 175) # light grey
        
        # Configure Footer
        footer = section.footer
        footer_p = footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_p.add_run("Page ")
        footer_run.font.name = 'Arial'
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(107, 114, 128)
        add_page_number_to_run(footer_run)

    # Configure Normal Style defaults
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = TEXT_COLOR
    
    # --- Title Page / Document Title ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_heading_paragraph(title_p, title, size_pt=24, color=PRIMARY_COLOR, bold=True, space_before=24, space_after=18)
    
    # Add a horizontal rule below title
    divider_p = doc.add_paragraph()
    divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider_p.add_run("—" * 40)
    divider_run.font.color.rgb = SECONDARY_COLOR
    divider_run.font.bold = True
    divider_p.paragraph_format.space_after = Pt(24)

    # --- Compile Sections ---
    for sec in sections:
        heading = sec.get("heading", "").strip()
        content = sec.get("content", "").strip()
        
        if not heading and not content:
            continue
            
        # Add Section Heading
        if heading:
            # Determine heading level based on style or sub-numbering (e.g. 1.1 vs 1.)
            is_sub = any(char.isdigit() for char in heading.split()[0]) and len(heading.split()[0].split('.')) > 2 if heading.split() else False
            if is_sub:
                p = doc.add_paragraph()
                style_heading_paragraph(p, heading, size_pt=14, color=SECONDARY_COLOR, bold=True, space_before=12, space_after=4)
            else:
                p = doc.add_paragraph()
                style_heading_paragraph(p, heading, size_pt=18, color=PRIMARY_COLOR, bold=True, space_before=18, space_after=8)
                
        # Parse content block by block (separated by double newlines)
        blocks = re.split(r'\n\n+', content)
        for block in blocks:
            block = block.strip()
            if not block:
                continue
                
            # Case 1: Markdown Table Detection
            if "|" in block and len([line for line in block.split('\n') if "|" in line]) >= 2:
                lines = [line.strip() for line in block.split('\n') if line.strip()]
                # Extract rows & cells
                table_data = []
                for line in lines:
                    # Skip separator row e.g., |---|---|
                    if re.match(r'^\|?[\s\-:|]+$', line):
                        continue
                    cells = [c.strip() for c in line.split('|')]
                    # If line starts and ends with '|', split leaves empty strings at start/end
                    if line.startswith('|'):
                        cells = cells[1:]
                    if line.endswith('|'):
                        cells = cells[:-1]
                    table_data.append(cells)
                
                if table_data:
                    num_cols = len(table_data[0])
                    num_rows = len(table_data)
                    word_table = doc.add_table(rows=num_rows, cols=num_cols)
                    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Compute dynamic widths (6.5 inches printable area)
                    col_width = Inches(6.5 / num_cols)
                    
                    for r_idx, row_cells in enumerate(table_data):
                        row = word_table.rows[r_idx]
                        
                        # Apply row heights/properties
                        trPr = row._tr.get_or_add_trPr()
                        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>')) # Avoid page break inside row
                        if r_idx == 0:
                            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>')) # Repeat header on new page
                        
                        for c_idx, cell_val in enumerate(row_cells):
                            # Handle uneven columns safely
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.width = col_width
                                cell.text = cell_val
                                
                                # Format paragraph in cell
                                cp = cell.paragraphs[0]
                                cp.paragraph_format.space_before = Pt(4)
                                cp.paragraph_format.space_after = Pt(4)
                                
                                # Style run
                                for run in cp.runs:
                                    run.font.name = 'Arial'
                                    run.font.size = Pt(10)
                                    if r_idx == 0:
                                        run.font.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255) # White text for headers
                                    else:
                                        run.font.color.rgb = TEXT_COLOR
                                
                                # Background colors
                                if r_idx == 0:
                                    set_cell_background(cell, SHADING_HEX_HEADER)
                                elif r_idx % 2 == 1:
                                    set_cell_background(cell, SHADING_HEX_ALT)
                    
                    # Add space after table
                    spacer = doc.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(8)
                    spacer.paragraph_format.space_after = Pt(4)
                    
            # Case 2: Callout Box Detection
            elif block.startswith('>') or block.startswith('! '):
                # Clean up quotes
                clean_lines = []
                for line in block.split('\n'):
                    line = line.strip()
                    if line.startswith('>'):
                        line = line[1:].strip()
                    elif line.startswith('!'):
                        line = line[1:].strip()
                    # Strip markdown block quotes alert tags like [!NOTE], [!IMPORTANT]
                    line = re.sub(r'^\[!(NOTE|IMPORTANT|WARNING|TIP|CAUTION)\]', '', line).strip()
                    clean_lines.append(line)
                clean_text = "\n".join(clean_lines).strip()
                
                # Render callout inside a 1x1 table
                callout_table = doc.add_table(rows=1, cols=1)
                callout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                cell = callout_table.cell(0, 0)
                cell.width = Inches(6.5)
                
                set_cell_background(cell, CALLOUT_BG_HEX)
                set_cell_left_border(cell, BORDER_COLOR_HEX, size_pt=3.5)
                
                cp = cell.paragraphs[0]
                cp.paragraph_format.left_indent = Inches(0.15)
                cp.paragraph_format.right_indent = Inches(0.15)
                cp.paragraph_format.space_before = Pt(8)
                cp.paragraph_format.space_after = Pt(8)
                
                run = cp.add_run(clean_text)
                run.font.name = 'Arial'
                run.font.italic = True
                run.font.size = Pt(10.5)
                run.font.color.rgb = PRIMARY_COLOR
                
                # Add spacing after callout
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_before = Pt(6)
                spacer.paragraph_format.space_after = Pt(2)
                
            # Case 3: List Bullet / Number Detection
            elif block.startswith('-') or block.startswith('*') or re.match(r'^\d+\.', block):
                lines = block.split('\n')
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    is_numbered = re.match(r'^(\d+)\.(.*)', line)
                    is_bullet = line.startswith('-') or line.startswith('*')
                    
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(4)
                    
                    if is_bullet:
                        p.paragraph_format.left_indent = Inches(0.25)
                        # Strip bullet symbol
                        clean_line = line[1:].strip()
                        # Use simple bullet run format
                        run_bullet = p.add_run("•  ")
                        run_bullet.font.name = 'Arial'
                        run_bullet.font.bold = True
                        run_bullet.font.color.rgb = SECONDARY_COLOR
                        
                        run_text = p.add_run(clean_line)
                        run_text.font.name = 'Arial'
                        run_text.font.size = Pt(11)
                        run_text.font.color.rgb = TEXT_COLOR
                        
                    elif is_numbered:
                        p.paragraph_format.left_indent = Inches(0.25)
                        num = is_numbered.group(1)
                        clean_line = is_numbered.group(2).strip()
                        
                        run_num = p.add_run(f"{num}.  ")
                        run_num.font.name = 'Arial'
                        run_num.font.bold = True
                        run_num.font.color.rgb = SECONDARY_COLOR
                        
                        run_text = p.add_run(clean_line)
                        run_text.font.name = 'Arial'
                        run_text.font.size = Pt(11)
                        run_text.font.color.rgb = TEXT_COLOR
                    else:
                        run_text = p.add_run(line)
                        run_text.font.name = 'Arial'
                        run_text.font.size = Pt(11)
                        run_text.font.color.rgb = TEXT_COLOR
                        
            # Case 4: Default Body Paragraph
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                # Check for inline bold formatting e.g., **text**
                parts = re.split(r'(\*\*.*?\*\*)', block)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        clean_part = part[2:-2]
                        run = p.add_run(clean_part)
                        run.font.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = TEXT_COLOR

    # --- Save Document ---
    filename = f"{session_id}_document.docx"
    filepath = os.path.join(settings.OUTPUT_DIR, filename)
    doc.save(filepath)
    return filepath
